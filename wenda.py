from pymongo import MongoClient
import docx
import re
import jieba

class get_save_search():
    # 建立数据库连接
    client = MongoClient('localhost', 27017)
    # 连接指定数据库
    db = client.py06db

    def getinfo(self):
        li = []
        file = docx.Document('words.docx')
        print("段落数："+str(len(file.paragraphs)))
        for i in file.paragraphs:
            details = i.text
            if details.strip().isdigit():
                continue
            else:
                details1 = re.sub(r'(^\d+| )','',details)
                li.append(details1.strip())
        return li

    def save_data(self,li):
        for item in li:
            try:
                if '答案' in item:
                    dict1 = {'question':li[li.index(item)-1],'result':item}
                    self.db.ask_answer.insert_one(dict1)
            except Exception as inde:
                continue


    def search_data(self,key_word):
        ls = []
        content = self.db.ask_answer.find()
        for item1 in content:
            question1 = item1['question']
            after_jieba = (','.join(jieba.cut(question1)))
            if key_word in after_jieba:
                # print(item1['question'],item1['result'])
                ls.append(item1)
        return ls

if __name__ == '__main__':
    AJ = get_save_search()
    # li = AJ.getinfo()
    #AJ.save_data(li)
    key_word = input('请输入关键字：')
    list1 = AJ.search_data(key_word)
    if list1 != []:
        for i in list1:
            print(i['question'], i['result'])
    else:
        print('搜索的关键字不存在！')






