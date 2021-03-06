# #读取docx中的文本代码示例
# import docx
# import re
#
# #获取文档
# file=docx.Document("word.docx")
# print("段落数:"+str(len(file.paragraphs))) #输出段落数
# file_word = docx.Document()
#
# #输出每一段的内容
# # for para in file.paragraphs:
# #     print(para.text)
#
# #输出段落编号及段落内容
# para_data = []
# for i in range(len(file.paragraphs)):
#     # for j in map(lambda x:x.split(' '),file.paragraphs[i].text):
#     para_single = file.paragraphs[i].text.split(' ')
#     while '' in para_single:  # 移除空格
#         para_single.remove('')
#     # para_data.append(para_single)
#     for data_number in range(len(para_single)):
#         data_num = re.findall(r"\d", para_single[data_number])
#         data_num = ''.join(data_num)
#         para_data.append(data_num + '    ')
# file_word.add_paragraph(para_data)
# file_word.save("word1.docx")





from pymongo import MongoClient as MC
import docx

#1、建立MongoDB 数据库连接
client = MC('localhost',27017)

#2、连接指定数据库
db = client.py06db
print(type(db.collection_names))
#获取当前指定数据库中的所有collection集合
print(db.collection_names())

#3、获取指定的collection集合
ask_answer = db.ask_answer

#获取数据并插入mongodb数据库
file = docx.Document('word.docx')
print("段落数："+str(len(file.paragraphs)))
for i in range(len(file.paragraphs)):
    details = file.paragraphs[i].text
    ask_answer.insert(details)